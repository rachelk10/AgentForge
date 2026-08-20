import io
import uuid
from typing import Any

from fastapi import HTTPException, status
from sqlalchemy import select

from app.models.document import Document, DocumentChunk
from app.rag.chunking import chunk_text, clean_text
from app.rag.embeddings import (
    EmbeddingProvider,
    OpenAIEmbeddingProvider,
)


class DocumentService:
    def __init__(
        self,
        db: Any,
        embedding_provider: EmbeddingProvider | None = None,
    ) -> None:
        self.db = db
        self.embedding_provider = embedding_provider

    @staticmethod
    def clean_text(text: str) -> str:
        return clean_text(text)

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 120, overlap: int = 30) -> list[str]:
        return chunk_text(text, chunk_size=chunk_size, overlap=overlap)

    @staticmethod
    def extract_text_from_bytes(filename: str, content: bytes) -> str:
        suffix = filename.lower()

        if suffix.endswith(".txt") or suffix.endswith(".md"):
            return content.decode("utf-8", errors="replace")

        if suffix.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except ImportError as exc:  # pragma: no cover
                raise ValueError("PDF support requires pypdf to be installed") from exc

            reader = PdfReader(io.BytesIO(content))
            pages: list[str] = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            return "\n\n".join(pages)

        if suffix.endswith(".docx"):
            try:
                from docx import Document
            except ImportError as exc:  # pragma: no cover
                raise ValueError("DOCX support requires python-docx to be installed") from exc

            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            return "\n\n".join(paragraphs)

        raise ValueError(f"Unsupported file type: {filename}")

    async def list_documents(self, agent_id: uuid.UUID) -> list[Document]:
        if self.db is None:
            return []

        result = await self.db.execute(
            select(Document).where(Document.agent_id == agent_id).order_by(Document.created_at.desc())
        )
        return list(result.scalars().all())

    async def get_document(self, agent_id: uuid.UUID, document_id: uuid.UUID) -> Document:
        if self.db is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

        result = await self.db.execute(
            select(Document).where(Document.id == document_id, Document.agent_id == agent_id)
        )
        document = result.scalar_one_or_none()
        if document is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
        return document

    async def delete_document(self, agent_id: uuid.UUID, document_id: uuid.UUID) -> None:
        document = await self.get_document(agent_id, document_id)
        await self.db.delete(document)
        await self.db.commit()

    async def upload_document(
        self,
        agent_id: uuid.UUID,
        filename: str,
        content: bytes,
        content_type: str | None = None,
    ) -> Document:
        if self.db is None:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Database not configured")

        extracted_text = self.extract_text_from_bytes(filename, content)
        cleaned_text = self.clean_text(extracted_text)
        chunks = self.chunk_text(cleaned_text)
        provider = self.embedding_provider or OpenAIEmbeddingProvider()
        embeddings = await provider.embed(chunks)
        if len(embeddings) != len(chunks):
            raise ValueError("Embedding provider returned an invalid result count")

        document = Document(
            agent_id=agent_id,
            filename=filename,
            content_type=content_type,
            file_size=len(content),
            status="processed",
            extracted_text=cleaned_text[:10000],
            metadata_={
                "source_type": filename.lower().rsplit(".", 1)[-1] if "." in filename else None,
                "chunk_count": len(chunks),
                "extracted_chars": len(cleaned_text),
                "embedding_model": getattr(provider, "model", provider.__class__.__name__),
                "embedding_dimensions": len(embeddings[0]) if embeddings else 0,
            },
        )
        self.db.add(document)
        await self.db.flush()

        for index, chunk in enumerate(chunks):
            self.db.add(
                DocumentChunk(
                    agent_id=agent_id,
                    document_id=document.id,
                    chunk_index=index,
                    content=chunk,
                    embedding=embeddings[index],
                    metadata_={
                        "filename": filename,
                        "source_type": filename.lower().rsplit(".", 1)[-1]
                        if "." in filename
                        else None,
                    },
                )
            )

        await self.db.commit()
        await self.db.refresh(document)
        return document
